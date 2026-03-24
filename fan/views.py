from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.db.models import Case, When, IntegerField
from django.http import JsonResponse, HttpResponse
import hashlib
import smtplib
from email.mime.text import MIMEText
from .models import CustomUser, Role, UserPreference, Match, League, Team,Sport
from .forms import RegistrationForm, UserPreferenceCreateForm, UserPreferenceUpdateForm, PasswordChangeForm
import csv
from io import StringIO
from django.contrib.auth.hashers import check_password
from django.http import FileResponse
from .reports import generate_user_report
from docx import Document
from io import BytesIO
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import League, Team, Match
from django.db.models import Q
import datetime



import os

def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.role = Role.objects.get_or_create(name="User")[0]
            user.save()
            login(request, user)
            return render(request, 'fan/register_success.html', {'message': 'Регистрация успешна!'})
        return render(request, 'fan/register.html', {'form': form, 'errors': form.errors})
    form = RegistrationForm()
    return render(request, 'fan/register.html', {'form': form})

def user_login(request):
    if request.method == 'POST':
        email = request.POST.get('email').strip().lower()  # Приводим email к нижнему регистру
        password = request.POST.get('password')
        print(f"Попытка входа: email={email}, password={password}")

        # Проверяем, существует ли пользователь в базе
        try:
            user = CustomUser.objects.get(email=email)
            print(f"Пользователь найден: {user.email}, is_active={user.is_active}")
            print(f"Хешированный пароль в базе: {user.password}")

            # Проверяем пароль вручную
            if check_password(password, user.password):
                print("Пароль совпадает вручную")
            else:
                print("Пароль НЕ совпадает вручную")
        except CustomUser.DoesNotExist:
            print("Пользователь с таким email не найден в базе")
            return render(request, 'fan/login.html', {
                'error': 'Пользователь с таким email не найден.'
            })

        # Пробуем аутентифицировать
        user = authenticate(request, username=email, password=password)
        if user is not None:
            print(f"Аутентификация успешна: {user.email}")
            if user.is_active:
                login(request, user)
                print("Успешный вход")
                return redirect('profile')
            else:
                print("Пользователь не активен")
                return render(request, 'fan/login.html', {
                    'error': 'Ваш аккаунт не активен.'
                })
        else:
            print("Ошибка: пользователь не аутентифицирован")
            return render(request, 'fan/login.html', {
                'error': 'Неверный email или пароль.'
            })
    return render(request, 'fan/login.html')

@login_required
def user_logout(request):
    logout(request)
    return redirect('login')

@login_required
def profile(request):
    if request.method == 'POST':
        if 'change_password' in request.POST:
            password_form = PasswordChangeForm(request.user, request.POST)
            if password_form.is_valid():
                user = password_form.save()
                update_session_auth_hash(request, user)
                return render(request, 'fan/profile.html', {
                    'message': 'Пароль успешно изменён!',
                    'preferences': UserPreference.objects.filter(user=request.user),
                    'password_form': PasswordChangeForm(request.user)
                })
        elif 'delete_preference' in request.POST:
            preference_id = request.POST.get('preference_id')
            preference = get_object_or_404(UserPreference, pk=preference_id, user=request.user)
            preference.delete()
            if request.session.get('active_preference_id') == str(preference.id):
                del request.session['active_preference_id']
            return redirect('profile')
        elif 'update_colors' in request.POST:
            preference_id = request.POST.get('preference_id')
            preference = get_object_or_404(UserPreference, pk=preference_id, user=request.user)
            new_color_left = request.POST.get('color_left')
            new_color_right = request.POST.get('color_right')
            if new_color_left and new_color_right:
                preference.custom_color_left = new_color_left
                preference.custom_color_right = new_color_right
                preference.save()
            return redirect('profile')
        elif 'select_preference' in request.POST:
            preference_id = request.POST.get('preference_id')
            request.session['active_preference_id'] = preference_id
            return redirect('profile')

    preferences = UserPreference.objects.filter(user=request.user)
    active_preference_id = request.session.get('active_preference_id')
    active_preference = preferences.filter(id=active_preference_id).first() if active_preference_id else None

    first_preference_color_left = '#ffffff'
    first_preference_color_right = '#ffffff'
    if active_preference:
        first_preference_color_left = active_preference.custom_color_left or active_preference.team.color_left
        first_preference_color_right = active_preference.custom_color_right or active_preference.team.color_right

    team_colors = {pref.id: {
        'color_left': pref.custom_color_left or pref.team.color_left,
        'color_right': pref.custom_color_right or pref.team.color_right
    } for pref in preferences}

    password_form = PasswordChangeForm(request.user)
    return render(request, 'fan/profile.html', {
        'user': request.user,
        'preferences': preferences,
        'team_colors': team_colors,
        'first_preference_color_left': first_preference_color_left,
        'first_preference_color_right': first_preference_color_right,
        'password_form': password_form,
        'active_preference_id': active_preference_id
    })

@login_required
def schedule(request):
    preferences = UserPreference.objects.filter(user=request.user)
    active_preference_id = request.session.get('active_preference_id')
    active_preference = preferences.filter(id=active_preference_id).first() if active_preference_id else None

    first_preference_color_left = '#ffffff'
    first_preference_color_right = '#ffffff'
    if active_preference:
        first_preference_color_left = active_preference.custom_color_left or active_preference.team.color_left
        first_preference_color_right = active_preference.custom_color_right or active_preference.team.color_right

    # Получаем ВСЕ матчи из базы данных
    sort_by = request.POST.get('sort_by', 'date')  # Используем POST вместо GET
    print(f"Request method: {request.method}, Sort by: {sort_by}")  # Отладочный вывод

    matches = Match.objects.all()  # Все матчи

    # Сортировка
    if sort_by == 'preference':
        user_teams = [pref.team.id for pref in preferences]
        matches = matches.order_by(
            Case(
                *[When(team_id=team_id, then=pos) for pos, team_id in enumerate(user_teams)],
                default=len(user_teams),
                output_field=IntegerField()
            )
        )
    elif sort_by == 'sport':
        matches = matches.order_by('team__league__sport__name')
    elif sort_by == 'league':
        matches = matches.order_by('team__league__name')
    elif sort_by == 'date':
        matches = matches.order_by('date')

    return render(request, 'fan/schedule.html', {
        'matches': matches,
        'preferences': preferences,
        'first_preference_color_left': first_preference_color_left,
        'first_preference_color_right': first_preference_color_right,
        'sort_by': sort_by
    })

@login_required
def user_preference_create(request):
    if request.method == 'POST':
        form = UserPreferenceCreateForm(request.user, request.POST)
        if form.is_valid():
            preference = form.save(commit=False)
            preference.user = request.user
            preference.save()
            request.session['active_preference_id'] = str(preference.id)
            return redirect('profile')
        return render(request, 'fan/user_preference_create.html', {'form': form, 'errors': form.errors})
    form = UserPreferenceCreateForm(request.user)
    return render(request, 'fan/user_preference_create.html', {'form': form})

@login_required
def user_preference_list(request):
    preferences = UserPreference.objects.filter(user=request.user)
    return render(request, 'fan/user_preference_list.html', {'preferences': preferences})

@login_required
def user_preference_update(request, pk):
    preference = get_object_or_404(UserPreference, pk=pk, user=request.user)
    if request.method == 'POST':
        form = UserPreferenceUpdateForm(request.user, request.POST, instance=preference)
        if form.is_valid():
            form.save()
            return redirect('profile')
        return render(request, 'fan/user_preference_update.html', {'form': form, 'errors': form.errors})
    form = UserPreferenceUpdateForm(request.user, instance=preference)
    return render(request, 'fan/user_preference_update.html', {'form': form})

@login_required
def user_preference_delete(request, pk):
    preference = get_object_or_404(UserPreference, pk=pk, user=request.user)
    if request.method == 'POST':
        preference.delete()
        if request.session.get('active_preference_id') == str(pk):
            del request.session['active_preference_id']
        return redirect('profile')
    return render(request, 'fan/user_preference_delete.html', {'preference': preference})

def forgot_password(request):
    if request.method == 'POST':
        email = request.POST.get('email').strip().lower()
        print(f"Запрос восстановления пароля для email: {email}")
        try:
            user = CustomUser.objects.get(email=email)
            print(f"Пользователь найден: {user.email}")
            new_password = CustomUser.objects.make_random_password(length=12)
            user.set_password(new_password)
            user.save()
            print(f"Новый пароль установлен: {new_password}")

            msg = MIMEText(f"Ваш новый пароль: {new_password}")
            msg['Subject'] = 'Восстановление пароля'
            msg['From'] = 'your_email@gmail.com'
            msg['To'] = email

            try:
                with smtplib.SMTP('smtp.gmail.com', 587) as server:
                    server.starttls()
                    server.login('your_email@gmail.com', 'your_app_password')
                    server.send_message(msg)
                    print("Письмо с новым паролем отправлено")
                return render(request, 'fan/forgot_password.html', {
                    'message': 'Новый пароль отправлен на ваш email!'
                })
            except Exception as e:
                print(f"Ошибка отправки email: {str(e)}")
                return render(request, 'fan/forgot_password.html', {
                    'error': 'Ошибка при отправке email. Попробуйте позже.'
                })
        except CustomUser.DoesNotExist:
            print("Пользователь с таким email не найден")
            return render(request, 'fan/forgot_password.html', {
                'error': 'Пользователь с таким email не найден.'
            })
    return render(request, 'fan/forgot_password.html')

def get_leagues(request):
    sport_id = request.GET.get('sport_id')
    leagues = League.objects.filter(sport_id=sport_id).values('id', 'name')
    return JsonResponse(list(leagues), safe=False)

def get_teams(request):
    league_id = request.GET.get('league_id')
    teams = Team.objects.filter(league_id=league_id).values('id', 'name')
    return JsonResponse(list(teams), safe=False)

@login_required
def export(request):
    if request.method == 'POST':
        export_type = request.POST.get('export_type', 'excel')
        export_fan_match = 'export_fan_match' in request.POST

        if not export_fan_match:
            return render(request, 'fan/export.html', {'error': 'Выберите данные для экспорта.'})

        matches = Match.objects.all()

        if export_type == 'excel':
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment; filename="fan_match.csv"'

            writer = csv.writer(response)
            writer.writerow(['ID', 'Team', 'Opponent', 'Date', 'Location', 'League', 'Sport'])
            for match in matches:
                writer.writerow([
                    match.id,
                    match.team.name,
                    match.opponent.name,
                    match.date,
                    match.location,
                    match.team.league.name,
                    match.team.league.sport.name
                ])
            return response

        elif export_type == 'html':
            output = StringIO()
            output.write('<html><head><meta charset="UTF-8"><title>Fan Match Export</title></head>')
            output.write('<body bgcolor="800000"><table align="center" cols="0" cellspacing="0">')
            output.write('<tr>')
            for header in ['ID', 'Team', 'Opponent', 'Date', 'Location', 'League', 'Sport']:
                output.write(f'<td><font face="Verdana" size="2" color="#ffffff"><p align="center"><b>{header}</b></font></td>')
            output.write('</tr>')
            for i, match in enumerate(matches, 1):
                bgcolor = '3399' if i % 2 == 0 else ''
                color = '#000000' if i % 2 == 0 else '#ffffff'
                output.write(f'<tr bgcolor="{bgcolor}">')
                row = [
                    match.id,
                    match.team.name,
                    match.opponent.name,
                    match.date,
                    match.location,
                    match.team.league.name,
                    match.team.league.sport.name
                ]
                for value in row:
                    output.write(f'<td><font face="Verdana" size="2" color="{color}"><p align="center">{value}</font></td>')
                output.write('</tr>')
            output.write('</table></body></html>')

            response = HttpResponse(content_type='text/html')
            response['Content-Disposition'] = 'attachment; filename="fan_match.html"'
            response.write(output.getvalue())
            return response

    return render(request, 'fan/export.html')

def download_league_report(request, league_id):
    league = get_object_or_404(League, id=league_id)
    teams = Team.objects.filter(league=league)
    matches = Match.objects.filter(team__league=league)

    doc = Document()
    doc.add_heading(f'Отчёт по лиге: {league.name}', level=1)
    doc.add_paragraph(f'Вид спорта: {league.sport.name}')
    doc.add_paragraph(f'Количество команд: {teams.count()}')

    doc.add_heading('Команды:', level=2)
    for team in teams:
        doc.add_paragraph(f'- {team.name} (Цвета: {team.color_left} / {team.color_right})')

    doc.add_heading('Матчи:', level=2)
    if matches.exists():
        for match in matches:
            doc.add_paragraph(f"{match.team.name} vs {match.opponent.name} — {match.date.strftime('%d.%m.%Y %H:%M')} ({match.location})")
    else:
        doc.add_paragraph("Нет зарегистрированных матчей.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=league_report_{league.name}.docx'
    return response

@login_required
def download_user_report(request):
    from .reports import generate_user_report  # если функция лежит в reports.py
    buffer = BytesIO()
    generate_user_report(buffer)  # исправлено: функция теперь получает BytesIO
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="user_report.docx"'
    return response
@login_required
def league_report_view(request):
    leagues = League.objects.all()
    return render(request, 'fan/league_report_list.html', {'leagues': leagues})

from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect, get_object_or_404

@login_required
def set_theme(request):
    if request.method == 'POST':
        preference_id = request.POST.get('preference_id')
        request.session['active_preference_id'] = preference_id
    return redirect(request.META.get('HTTP_REFERER', 'profile'))

from docx import Document
from io import BytesIO
from django.http import HttpResponse
@login_required
def league_report_view(request):
    user = request.user
    sports = Sport.objects.all()
    leagues = League.objects.prefetch_related('teams')
    preferences = UserPreference.objects.filter(user=user)
    active_preference = preferences.first()

    selected_sport_id = request.POST.get('sport', '')
    selected_league_ids = request.POST.getlist('leagues')
    selected_preference_ids = request.POST.getlist('preferences')

    if selected_sport_id:
        leagues = leagues.filter(sport_id=selected_sport_id)

    # Добавляем поле matches_list к каждой лиге
    for league in leagues:
        teams = league.teams.all()
        league.matches_list = Match.objects.filter(
            Q(team__in=teams) | Q(opponent__in=teams)
        ).order_by('date')

    preview_report = False

    # Если нажали "Показать превью отчета"
    if request.method == "POST" and "generate_preview" in request.POST:
        preview_report = True

    # Если нажали "Сохранить в Word"
    if request.method == "POST" and "save_word" in request.POST:
        document = Document()
        document.add_heading('Отчет по лигам', 0)

        for league in leagues:
            if str(league.id) not in selected_league_ids:
                continue
            document.add_heading(f'{league.name} ({league.sport.name})', level=1)

            document.add_heading('Команды', level=2)
            for team in league.teams.all():
                if active_preference and active_preference.team.id == team.id:
                    document.add_paragraph(f'{team.name} (активное предпочтение)')
                else:
                    document.add_paragraph(team.name)

            document.add_heading('Матчи', level=2)
            for match in league.matches_list:
                if (not selected_preference_ids or
                    str(match.team.id) in selected_preference_ids or
                    str(match.opponent.id) in selected_preference_ids):
                    document.add_paragraph(f'{match.team.name} vs {match.opponent.name} — {match.date.strftime("%d.%m.%Y %H:%M")} ({match.location})')

        # Формируем ответ с Word файлом
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=league_report_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        document.save(response)
        return response

    context = {
        'sports': sports,
        'leagues': leagues,
        'preferences': preferences,
        'active_preference': active_preference,
        'selected_sport_id': selected_sport_id,
        'selected_league_ids': selected_league_ids,
        'selected_preference_ids': selected_preference_ids,
        'preview_report': preview_report,
        'first_preference_color_left': getattr(active_preference, 'custom_color_left', '#ffffff'),
        'first_preference_color_right': getattr(active_preference, 'custom_color_right', '#ffffff'),
    }

    return render(request, 'fan/league_report.html', context)