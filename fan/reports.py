from io import BytesIO
from django.http import HttpResponse
from docx import Document
from .models import CustomUser, Sport, League, Team, Match

# --- Отчёт по пользователям ---
def generate_user_report():
    doc = Document()
    doc.add_heading('Отчёт по пользователям', level=1)
    users = CustomUser.objects.all().select_related('role')

    for user in users:
        doc.add_paragraph(f"{user.last_name} {user.first_name} — {user.email}")
        doc.add_paragraph(f"Роль: {user.role.name if user.role else '—'}")
        doc.add_paragraph(f"Активен: {'Да' if user.is_active else 'Нет'}")
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Отчёт по видам спорта ---
def generate_sport_report():
    doc = Document()
    doc.add_heading('Отчёт по видам спорта', level=1)
    sports = Sport.objects.all()

    for sport in sports:
        doc.add_paragraph(f"Вид спорта: {sport.name}")
        leagues = sport.leagues.all()
        for league in leagues:
            doc.add_paragraph(f"   • Лига: {league.name}")
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Отчёт по лигам ---
def generate_league_report(league_id=None):
    doc = Document()

    if league_id:
        league = League.objects.get(pk=league_id)
        doc.add_heading(f'Отчёт по лиге: {league.name}', level=1)
        teams = league.teams.all()
        for team in teams:
            doc.add_paragraph(f"Команда: {team.name} ({team.color_left} / {team.color_right})")
    else:
        doc.add_heading('Отчёт по всем лигам', level=1)
        leagues = League.objects.all()
        for league in leagues:
            doc.add_paragraph(f"Лига: {league.name} ({league.sport.name})")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Отчёт по командам ---
def generate_team_report():
    doc = Document()
    doc.add_heading('Отчёт по командам', level=1)
    teams = Team.objects.select_related('league', 'league__sport')

    for team in teams:
        doc.add_paragraph(f"{team.name} ({team.league.name}, {team.league.sport.name})")
        doc.add_paragraph(f"Цвета: {team.color_left} — {team.color_right}")
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Отчёт по матчам ---
def generate_match_report():
    doc = Document()
    doc.add_heading('Отчёт по матчам', level=1)
    matches = Match.objects.select_related('team', 'opponent')

    for match in matches:
        doc.add_paragraph(
            f"{match.team.name} vs {match.opponent.name} — {match.date.strftime('%d.%m.%Y %H:%M')} ({match.location})"
        )
    doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
